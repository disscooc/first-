import os, re, shutil, copy
from datetime import datetime
import docx
import config
from utils import safe_fn, new_path, handle_err, read_json, write_json, fmt_money

def _replace_all(doc, vars_map):
    for p in doc.paragraphs:
        _replace_para(p, vars_map)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_para(p, vars_map)
    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_para(p, vars_map)
        for p in section.footer.paragraphs:
            _replace_para(p, vars_map)
        for tbl in section.header.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_para(p, vars_map)
        for tbl in section.footer.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_para(p, vars_map)

def _replace_para(para, vars_map):
    full_text = para.text
    if not full_text:
        return
    for run in para.runs:
        for k, v in vars_map.items():
            if k in run.text:
                run.text = run.text.replace(k, str(v) if v else "")
    new_text = para.text
    for k, v in vars_map.items():
        if k in new_text:
            new_text = new_text.replace(k, str(v) if v else "")
    if new_text != para.text:
        if not para.runs:
            para.add_run(new_text)
        else:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""

def generate(unit, date, company, manager, phone,
             party_a_name, party_a_contact, party_a_phone, party_a_email, party_a_address, party_a_mailing_address, party_a_credit,
             party_a_bank, party_a_account,
             party_b_name, party_b_contact, party_b_phone, party_b_email, party_b_address, party_b_mailing_address, party_b_credit,
             party_b_bank, party_b_account,
             sign_place, sign_time, service_term, pay_days, invoice,
             total_amount, comp_items, stor_items, project=""):
    tpl = config.CONTRACT_TPL
    if not os.path.exists(tpl):
        _create_default_template(tpl)
    out = new_path(safe_fn(unit), "计算资源服务合同", ".docx")
    shutil.copy2(tpl, out)
    doc = docx.Document(out)

    # ========== 调试日志 ==========
    def _log_ct(msg):
        print(f"[CONTRACT_DBG] {msg}", flush=True)
    _log_ct("=" * 60)
    _log_ct(f"total_amount={total_amount} comp({len(comp_items)}) stor({len(stor_items)})")
    for i, it in enumerate(comp_items):
        _log_ct(f"  COMP[{i}]: rtype={it.rtype} name={it.name} spec={str(it.spec)[:30]}")
        _log_ct(f"    price={it.price} qty={it.quantity} amount={it.amount} orig_price={it.orig_price} disc={it.discount}")
    for i, it in enumerate(stor_items):
        _log_ct(f"  STOR[{i}]: rtype={it.rtype} name={it.name} spec={str(it.spec)[:30]}")
        _log_ct(f"    price={it.price} qty={it.quantity} amount={it.amount} free_quota={getattr(it,'free_quota',False)}")
    # ==============================

    # 输入的价格均为含税价，不含税价 = 含税价 / (1+税率)
    tax_total = round(total_amount, 2)
    no_tax = round(total_amount / config.TAX, 2)
    resource_list = ""
    idx = 1
    for it in comp_items:
        name = it.rtype + " - " + it.name
        if it.spec:
            name += " (" + it.spec + ")"
        qty_text = "" if getattr(it, "quote_mode", "") == "total" else "，数量 " + str(it.quantity or "")
        resource_list += str(idx) + ". " + name + "，单价 " + str(it.price) + " 元" + qty_text + "，金额 " + fmt_money(it.amount or 0) + "\n"
        _log_ct(f"  RESOURCE_LIST[{idx-1}]: {name}  price={it.price}  qty={it.quantity}  amount={it.amount}")
        idx += 1
    for it in stor_items:
        name = "存储 - " + it.name
        if it.spec:
            name += " (" + it.spec + ")"
        resource_list += str(idx) + ". " + name + "，单价 " + str(it.price) + " 元，数量 " + str(it.quantity or '') + "，金额 " + fmt_money(it.amount or 0) + "\n"
        _log_ct(f"  RESOURCE_LIST[{idx-1}]: {name}  price={it.price}  qty={it.quantity}  amount={it.amount}")
        idx += 1
    vars_map = {
        "{{甲方单位}}": party_a_name,
        "{{甲方联系人}}": party_a_contact,
        "{{甲方电话}}": party_a_phone,
        "{{甲方邮箱}}": party_a_email,
        "{{甲方地址}}": party_a_address,
        "{{甲方通信地址}}": party_a_mailing_address or party_a_address,
        "{{甲方信用代码}}": party_a_credit,
        "{{甲方开户行}}": party_a_bank,
        "{{甲方账号}}": party_a_account,
        "{{乙方单位}}": party_b_name or config.DEF_CONTRACT.get("party_b_name", company),
        "{{乙方联系人}}": party_b_contact or manager,
        "{{乙方电话}}": party_b_phone or phone,
        "{{乙方邮箱}}": party_b_email,
        "{{乙方地址}}": party_b_address or config.DEF_CONTRACT.get("party_b_addr", ""),
        "{{乙方通信地址}}": party_b_mailing_address or party_b_address or config.DEF_CONTRACT.get("party_b_mailing_address", config.DEF_CONTRACT.get("party_b_addr", "")),
        "{{乙方信用代码}}": party_b_credit or config.DEF_CONTRACT.get("party_b_credit", ""),
        "{{乙方开户行}}": party_b_bank or config.DEF_CONTRACT.get("party_b_bank", ""),
        "{{乙方账号}}": party_b_account or config.DEF_CONTRACT.get("party_b_account", ""),
        "{{签约地点}}": sign_place or config.DEF_CONTRACT.get("sign_place", ""),
        "{{签约日期}}": _year_month(sign_time or date),
        "{{含税总价}}": fmt_money(tax_total),
        "{{不含税总价}}": fmt_money(no_tax),
        "{{税率}}": "6%",
        "{{服务期限}}": service_term or config.DEF_CONTRACT.get("service_term", "1年"),
        "{{付款天数}}": pay_days or config.DEF_CONTRACT.get("pay_days", "30"),
        "{{发票类型}}": invoice or config.DEF_CONTRACT.get("invoice", "6%增值税普通发票"),
        "{{资源清单}}": resource_list.strip(),
        "{{客户单位}}": unit,
        "{{报价单位}}": company,
        "{{项目名称}}": project,
        "{{报价日期}}": date,
        "{{合同日期}}": date,
        "{{甲方}}": party_a_name,
        "{{乙方}}": party_b_name or config.DEF_CONTRACT.get("party_b_name", company),
    }
    _replace_all(doc, vars_map)
    _fill_real_contract_layout(
        doc, unit, date, company, manager, phone,
        party_a_name, party_a_contact, party_a_phone, party_a_email, party_a_address, party_a_mailing_address, party_a_credit,
        party_a_bank, party_a_account,
        party_b_name, party_b_contact, party_b_phone, party_b_email, party_b_address, party_b_mailing_address, party_b_credit,
        party_b_bank, party_b_account,
        sign_place, sign_time, service_term, pay_days, invoice,
        tax_total, no_tax, comp_items, stor_items, project
    )
    doc.save(out)
    _save_customer(unit, {
        # 公司信息
        "party_a_name": party_a_name,
        "party_a_address": party_a_address,
        "party_a_credit": party_a_credit,
        "party_a_bank": party_a_bank,
        "party_a_account": party_a_account,
        # 联系人信息（联系人/通讯地址/电话/邮箱）
        "party_a_contact": party_a_contact,
        "party_a_contact_mailing_address": party_a_mailing_address,
        "party_a_contact_phone": party_a_phone,
        "party_a_contact_email": party_a_email,
    })
    return out

def _fill_real_contract_layout(doc, unit, date, company, manager, phone,
                               party_a_name, party_a_contact, party_a_phone, party_a_email, party_a_address, party_a_mailing_address, party_a_credit,
                               party_a_bank, party_a_account,
                               party_b_name, party_b_contact, party_b_phone, party_b_email, party_b_address, party_b_mailing_address, party_b_credit,
                               party_b_bank, party_b_account,
                               sign_place, sign_time, service_term, pay_days, invoice,
                               tax_total, no_tax, comp_items, stor_items, project):
    paragraphs = doc.paragraphs
    party_section = None
    for para in paragraphs:
        text = para.text
        stripped = text.strip()
        # 首页：3个普通空格 "甲  方：北京理工大学"
        if stripped.startswith("甲  方："):
            _set_party_line_para(para, "甲  方：   ", party_a_name or unit)
        elif stripped.startswith("乙  方："):
            _set_party_line_para(para, "乙  方：   ", party_b_name or company)
        # 信息页：5个普通空格 "甲    方：北京理工大学"（需设置 party_section）
        elif stripped.startswith("甲    方："):
            _set_party_line_para(para, "甲    方：", party_a_name or unit)
            party_section = "A"
        elif stripped.startswith("乙    方："):
            _set_party_line_para(para, "乙    方：", party_b_name or company)
            party_section = "B"
        # 签字页：无空格 "甲方："
        elif stripped.startswith("甲方："):
            _set_para_text(para, f"甲方：{party_a_name or unit}")
        elif stripped.startswith("乙方："):
            _set_para_text(para, f"乙方：{party_b_name or company}")
        elif stripped.startswith("签订地点："):
            _set_para_text(para, f"签订地点： {sign_place or '北京'}")
        elif stripped.startswith("签订时间："):
            # 首页只写到月份
            _set_para_text(para, f"签订时间：{_year_month(sign_time or date)}")
        elif stripped.startswith("签字日期："):
            # 签字日期留空，由双方手工填写
            _set_para_text(para, "签字日期：       年    月    日")
        elif stripped.startswith("注册地址、电话："):
            value = party_a_address if party_section == "A" else party_b_address
            _set_para_text(para, f"注册地址、电话：{value}")
        elif stripped.startswith("统一社会信用代码："):
            value = party_a_credit if party_section == "A" else party_b_credit
            _set_para_text(para, f"统一社会信用代码：{value}")
        elif stripped.startswith("账号名称："):
            value = (party_a_name or unit) if party_section == "A" else (party_b_name or company)
            _set_para_text(para, f"账号名称：{value}")
        elif stripped.startswith("开户银行："):
            value = party_a_bank if party_section == "A" else party_b_bank
            _set_para_text(para, f"开户银行：{value}")
        elif stripped.startswith("开户账号："):
            value = party_a_account if party_section == "A" else party_b_account
            _set_para_text(para, f"开户账号：{value}")
        elif stripped.startswith("联 系 人："):
            value = party_a_contact if party_section == "A" else (party_b_contact or manager)
            _set_para_text(para, f"联 系 人：{value}")
        elif stripped.startswith("通信地址："):
            value = (party_a_mailing_address or party_a_address) if party_section == "A" else (party_b_mailing_address or party_b_address)
            _set_para_text(para, f"通信地址：{value}")
        elif stripped.startswith("电    话："):
            value = party_a_phone if party_section == "A" else (party_b_phone or phone)
            _set_para_text(para, f"电    话：{value}")
        elif "甲方指定接收对账函的邮箱为：" in stripped:
            import re
            new_text = re.sub(r"甲方指定接收对账函的邮箱为：[^ ]+", 
                              f"甲方指定接收对账函的邮箱为：{party_a_email}", 
                              stripped)
            _set_para_text(para, new_text)
        elif stripped.startswith("1.1项目名称："):
            _set_para_text(para, f"1.1项目名称：{project or '计算资源服务项目'}")
        elif stripped.startswith("1.3项目服务期限："):
            _set_para_text(para, f"1.3项目服务期限：自实际充值之日起 {service_term or '1年'}。服务期限届满或核时（1核时即可使用1颗核心1小时）用尽合同自动终止。")
        elif stripped.startswith("甲方在本合同生效后"):
            # 金额格式与表格汇总行一致：含税总价 + 大写
            _set_para_text(para, f"甲方在本合同生效后【{pay_days or '30'}】日内，将合同所有费用{_money_with_cn(tax_total)}支付给乙方，在本合同有效期内，甲方可根据本合同约定的资源数量使用乙方所提供的计算服务（如未约定资源数量，则甲方按双方约定的资源单价使用乙方所提供的计算服务，按使用量实时扣减，总费用扣完为止）；")
    _fill_real_contract_table(doc, comp_items, stor_items, tax_total)

def _fill_real_contract_table(doc, comp_items, stor_items, tax_total):
    if len(doc.tables) < 2:
        return
    table = doc.tables[1]
    free_tb = sum((item.quantity or 0) for item in stor_items if getattr(item, "free_quota", False))
    if table.rows:
        _set_cell_with_underline(table.rows[0].cells[0], f"A. 计算资源（含赠送 {free_tb:g} TB存储）\n共享队列按需计费", underline_part=f"{free_tb:g}")

    comp = list(comp_items)
    comp_start, comp_end = _section_data_bounds(table, "A.", "B.")
    if comp_start is not None:
        comp_count = max(1, len(comp))
        _resize_table_block(table, comp_start, comp_end, comp_count)

    stor = list(stor_items)
    stor_start, stor_end = _section_data_bounds(table, "B.", "C.")
    if stor_start is not None:
        stor_count = max(1, len(stor))
        _resize_table_block(table, stor_start, stor_end, stor_count)

    if comp_start is not None:
        for index in range(len(comp)):
            row = table.rows[comp_start + index]
            item = comp[index]
            _set_cell_text(row.cells[0], item.name)
            _set_cell_text(row.cells[1], item.spec)
            _set_cell_text(row.cells[2], _fmt_price(item.price))
            _set_cell_text(row.cells[3], "")
            # 调试日志
            print(f"[CONTRACT_TBL_DBG] COMP[{index}] write_row={comp_start+index}  "
                  f"col0(name)={item.name}  col1(spec)={str(item.spec)[:20]}  "
                  f"col2(price)={_fmt_price(item.price)}", flush=True)
        if comp:
            _set_vertical_merged_cell_text(table, comp_start, comp_start + len(comp) - 1, 3, fmt_money(tax_total))

    if stor_start is not None:
        for index in range(len(stor)):
            row = table.rows[stor_start + index]
            item = stor[index]
            # 名称
            _set_cell_text(row.cells[0], item.name)
            # 技术规格：使用 List Price 中匹配到的规格；没有规格时再回退到名称
            stor_spec = item.spec if item.spec else (item.name if item.name else "分布式文件存储")
            _set_cell_text(row.cells[1], stor_spec)
            # 单价：折后价
            discounted_price = round((item.orig_price or item.price or 0) * (item.discount or 1), 4)
            if discounted_price > 0:
                _set_cell_text(row.cells[2], _fmt_price(discounted_price))
            else:
                _set_cell_text(row.cells[2], "0")
            # 总费用：显示实际金额
            _set_cell_text(row.cells[3], fmt_money(item.amount or 0))
            # 调试日志
            print(f"[CONTRACT_TBL_DBG] STOR[{index}] write_row={stor_start+index}  "
                  f"col0(name)={item.name}  col1(spec)={stor_spec}  "
                  f"col2(price)={_fmt_price(discounted_price)}  col3(amount)={fmt_money(item.amount or 0)}", flush=True)
        if not comp and stor:
            _set_vertical_merged_cell_text(table, stor_start, stor_start + len(stor) - 1, 3, fmt_money(tax_total))

    no_tax = round(tax_total / config.TAX, 2)
    tax = round(tax_total - no_tax, 2)
    for row in table.rows:
        label = row.cells[0].text.strip()
        # 注意："不含税" 必须优先于 "含税" 判断，否则 "不含税金额总计" 会误匹配到 "含税"
        if "不含税" in label and ("合计" in label or "总计" in label):
            _set_summary_row(row, _money_with_cn(no_tax))
        elif "含税" in label and ("合计" in label or "总计" in label):
            _set_summary_row(row, _money_with_cn(tax_total))
        elif "税额" in label and ("合计" in label or "总计" in label):
            _set_summary_row(row, _money_with_cn(tax, prefix=False))

def _set_para_text(para, text):
    """设置段落文本，保留原有字体"""
    if not para.runs:
        para.add_run(str(text))
        return
    # 保留第一个run的字体，只改文本
    para.runs[0].text = str(text)
    for run in para.runs[1:]:
        run.text = ""

def _set_party_line_para(para, label_text, name_value):
    """设置甲/乙方名称段落，首页四号14pt，信息页小四12pt，名称带下划线"""

    # 判断是首页还是信息页：首页标签 "甲  方：   "（3空格），信息页 "甲    方："（5空格）
    is_first_page = label_text.startswith("甲  方") or label_text.startswith("乙  方")

    # 清空所有 run 并重建
    for run in para.runs:
        run.text = ""
    while len(para.runs) < 2:
        para.add_run("")

    font_name = "仿宋"
    font_size = 14 if is_first_page else 12  # 四号14pt，小四12pt

    # 标签（无下划线，首页加粗）
    para.runs[0].text = label_text
    para.runs[0].underline = False
    para.runs[0].bold = bool(is_first_page)
    _set_run_font(para.runs[0], font_name, font_size)

    # 名称（带下划线，首页加粗）
    para.runs[1].text = str(name_value) if name_value else ""
    para.runs[1].underline = True
    para.runs[1].bold = bool(is_first_page)
    _set_run_font(para.runs[1], font_name, font_size)

    # 清空多余runs
    for run in para.runs[2:]:
        run.text = ""


def _copy_run_font(src_run, dst_run):
    """从src_run复制字体属性到dst_run（含东亚字体）"""
    if src_run is None or dst_run is None:
        return
    try:
        src = src_run.font
        dst = dst_run.font
        if src.name:
            dst.name = src.name
        if src.size:
            dst.size = src.size
        if src.bold is not None:
            dst.bold = src.bold
        if src.italic is not None:
            dst.italic = src.italic
        if src.color and src.color.rgb:
            dst.color.rgb = src.color.rgb
        # 复制东亚字体（python-docx font.name 只设西文，CJK 需走 XML）
        _copy_east_asian_font(src_run, dst_run)
    except Exception:
        pass  # 字体复制失败时保持默认


def _copy_east_asian_font(src_run, dst_run):
    """从 src_run 复制 w:eastAsia 字体名到 dst_run"""
    from lxml import etree
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    src_rpr = src_run._element.find(".//w:rPr", nsmap)
    dst_rpr = dst_run._element.find(".//w:rPr", nsmap)
    if src_rpr is None or dst_rpr is None:
        return
    src_rfonts = src_rpr.find("w:rFonts", nsmap)
    if src_rfonts is None:
        return
    east = src_rfonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")
    if not east:
        return
    dst_rfonts = dst_rpr.find("w:rFonts", nsmap)
    if dst_rfonts is None:
        dst_rfonts = etree.SubElement(dst_rpr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
    dst_rfonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", east)


def _set_run_font(run, font_name="仿宋", font_size_pt=None, east_asian=None):
    """显式设置 run 字体（西文+东亚），并可选字号"""
    from lxml import etree
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    rpr = run._element.find(".//w:rPr", nsmap)
    if rpr is None:
        rpr = etree.SubElement(run._element, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
    rfonts = rpr.find("w:rFonts", nsmap)
    if rfonts is None:
        rfonts = etree.SubElement(rpr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
    # 西文字体
    if font_name:
        rfonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii", font_name)
        rfonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi", font_name)
    # 东亚字体（缺省与西文一致）
    ea = east_asian or font_name
    if ea:
        rfonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", ea)
    # python-docx 层面也设一次
    if font_name:
        run.font.name = font_name
    # 字号
    if font_size_pt is not None:
        from docx.shared import Pt
        run.font.size = Pt(font_size_pt)
        # 同时设置 szCs（复杂脚本字号，对中文生效）
        sz = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz")
        szCs = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}szCs")
        sz_val = str(int(font_size_pt * 2))  # Pt → half-points
        if sz is None:
            sz = etree.SubElement(rpr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz")
        sz.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", sz_val)
        if szCs is None:
            szCs = etree.SubElement(rpr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}szCs")
        szCs.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", sz_val)


def _clear_all_runs(para):
    """清空段落所有runs的文本（保留以便兼容）"""
    for run in para.runs:
        run.text = ""

def _set_cell_text(cell, text):
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    _set_para_text(para, text)
    for extra_para in cell.paragraphs[1:]:
        _set_para_text(extra_para, "")


def _set_cell_with_underline(cell, full_text, underline_part):
    """设置单元格文本，其中 underline_part 部分带下划线"""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    # 清空
    for run in para.runs:
        run.text = ""
    # 找到下划线部分的起始位置
    idx = full_text.find(underline_part)
    if idx >= 0:
        # 前段（无下划线）
        if idx > 0:
            if len(para.runs) < 1:
                para.add_run("")
            para.runs[0].text = full_text[:idx]
            para.runs[0].underline = False
        else:
            if para.runs:
                para.runs[0].text = ""
        # 中段（有下划线）
        while len(para.runs) < 2:
            para.add_run("")
        para.runs[1].text = underline_part
        para.runs[1].underline = True
        # 后段（无下划线）
        after = full_text[idx + len(underline_part):]
        if after:
            while len(para.runs) < 3:
                para.add_run("")
            para.runs[2].text = after
            para.runs[2].underline = False
        # 清空多余
        for run in para.runs[3:]:
            run.text = ""
    else:
        _set_para_text(para, full_text)
    for extra_para in cell.paragraphs[1:]:
        _set_para_text(extra_para, "")

def _section_data_bounds(table, section_prefix, next_prefix):
    section_idx = None
    next_idx = None
    for index, row in enumerate(table.rows):
        first = row.cells[0].text.strip()
        if section_idx is None and first.startswith(section_prefix):
            section_idx = index
            continue
        if section_idx is not None and first.startswith(next_prefix):
            next_idx = index
            break
    if section_idx is None or next_idx is None or section_idx + 2 >= next_idx:
        return None, None
    return section_idx + 2, next_idx

def _resize_table_block(table, start_idx, end_idx, desired_count):
    current_count = max(0, end_idx - start_idx)
    if current_count <= 0:
        return
    while current_count < desired_count:
        _clone_table_row_before(table, start_idx, end_idx)
        current_count += 1
        end_idx += 1
    while current_count > desired_count:
        _delete_table_row(table, start_idx + current_count - 1)
        current_count -= 1

def _delete_table_row(table, row_idx):
    tr = table.rows[row_idx]._tr
    tr.getparent().remove(tr)

def _clear_row(row):
    for cell in row.cells:
        _set_cell_text(cell, "")

def _clone_table_row_before(table, template_idx, before_idx):
    new_tr = copy.deepcopy(table.rows[template_idx]._tr)
    table.rows[before_idx]._tr.addprevious(new_tr)

def _fmt_price(value):
    try:
        value = float(value or 0)
    except (TypeError, ValueError):
        value = 0
    return f"{value:g}"

def _set_summary_row(row, text):
    """设置汇总行：cells[1:]合并后填入金额"""
    cells = row.cells
    if len(cells) >= 3:
        # 合并第2个到最后一个单元格
        try:
            cells[1].merge(cells[-1])
        except Exception:
            pass
    # 填入汇总金额
    target = cells[1] if len(cells) > 1 else cells[0]
    _set_cell_text(target, text)


def _set_vertical_merged_cell_text(table, start_row, end_row, col_idx, text):
    if start_row is None or end_row is None or end_row < start_row:
        return
    try:
        target = table.rows[start_row].cells[col_idx]
        if end_row > start_row:
            target = target.merge(table.rows[end_row].cells[col_idx])
        _set_cell_text(target, text)
    except Exception:
        _set_cell_text(table.rows[start_row].cells[col_idx], text)
        for row_idx in range(start_row + 1, end_row + 1):
            _set_cell_text(table.rows[row_idx].cells[col_idx], "")

def _money_with_cn(value, prefix=True):
    amount = round(float(value or 0), 2)
    cn = _rmb_upper(amount)
    head = "人民币" if prefix else ""
    return f"【{fmt_money(amount)}】元（大写：{head}{cn}）"

def _rmb_upper(value):
    value = round(float(value or 0), 2)
    integer = int(value)
    fraction = int(round((value - integer) * 100))
    digits = "零壹贰叁肆伍陆柒捌玖"
    units = ["", "拾", "佰", "仟"]
    groups = ["", "万", "亿", "兆"]

    def group_to_cn(number):
        result = ""
        zero = False
        for pos in range(3, -1, -1):
            base = 10 ** pos
            digit = number // base
            number %= base
            if digit:
                if zero:
                    result += "零"
                    zero = False
                result += digits[digit] + units[pos]
            elif result:
                zero = True
        return result

    if integer == 0:
        int_text = "零元"
    else:
        parts = []
        group_index = 0
        while integer:
            group = integer % 10000
            if group:
                parts.append(group_to_cn(group) + groups[group_index])
            integer //= 10000
            group_index += 1
        int_text = "零".join(reversed(parts)).replace("零零", "零") + "元"

    jiao = fraction // 10
    fen = fraction % 10
    if jiao == 0 and fen == 0:
        return int_text + "整"
    frac_text = ""
    if jiao:
        frac_text += digits[jiao] + "角"
    elif fen:
        frac_text += "零"
    if fen:
        frac_text += digits[fen] + "分"
    return int_text + frac_text


def _year_month(date_str):
    """从日期字符串中提取年月，如 '2026.06.04' → '2026.06'"""
    if not date_str:
        return datetime.now().strftime("%Y.%m")
    # 去掉日部分，保留年月
    parts = str(date_str).replace("-", ".").replace("/", ".").split(".")
    if len(parts) >= 2:
        return f"{parts[0]}.{parts[1]}"
    return str(date_str)


def _save_customer(unit_name, info):
    if not unit_name:
        return
    data = read_json(config.F_PRIVATE_CUSTOMERS, [])
    if isinstance(data, dict):
        data = list(data.values())
    if not isinstance(data, list):
        data = []
    record = dict(info)
    record["party_a_name"] = info.get("party_a_name") or unit_name
    replaced = False
    for index, item in enumerate(data):
        name = item.get("party_a_name") or item.get("PartyAName") or ""
        if name == unit_name:
            data[index] = record
            replaced = True
            break
    if not replaced:
        data.append(record)
    write_json(config.F_PRIVATE_CUSTOMERS, data)

def load_customer(unit_name):
    data = read_json(config.F_CTRACT, {})
    if isinstance(data, list):
        data = {}
    saved = data.get(unit_name, {}) if unit_name else {}
    library = _ensure_university_contracts()
    library_info = {}
    if unit_name in library:
        library_info = library.get(unit_name, {})
    else:
        for name, info in library.items():
            if unit_name and (unit_name in name or name in unit_name):
                library_info = info
                break
    private_info = _find_private_customer(unit_name)
    if saved:
        merged = dict(library_info)
        merged.update({key: value for key, value in saved.items() if value})
        merged.update({key: value for key, value in private_info.items() if value})
        return merged
    if private_info:
        merged = dict(library_info)
        merged.update({key: value for key, value in private_info.items() if value})
        return merged
    return library_info

def ensure_university_contracts():
    return _ensure_university_contracts()

def _normalize_private_customer(raw):
    if not isinstance(raw, dict):
        return {}
    party_a_name = raw.get("party_a_name") or raw.get("PartyAName") or ""
    party_a_address = raw.get("party_a_address") or raw.get("PartyAAddressPhone") or raw.get("PartyAAddress") or ""
    party_b_address = raw.get("party_b_address") or raw.get("PartyBAddressPhone") or raw.get("PartyBAddress") or ""
    return {
        "party_a_name": party_a_name,
        "party_a_contact": raw.get("party_a_contact") or raw.get("PartyAContact") or "",
        "party_a_phone": raw.get("party_a_phone") or raw.get("PartyAPhone") or "",
        "party_a_address": party_a_address,
        "party_a_credit": raw.get("party_a_credit") or raw.get("PartyACreditCode") or "",
        "party_a_bank": raw.get("party_a_bank") or raw.get("PartyABank") or "",
        "party_a_account": raw.get("party_a_account") or raw.get("PartyAAccountNo") or "",
        "party_b_name": raw.get("party_b_name") or raw.get("PartyBName") or "",
        "party_b_contact": raw.get("party_b_contact") or raw.get("PartyBContact") or "",
        "party_b_phone": raw.get("party_b_phone") or raw.get("PartyBPhone") or "",
        "party_b_address": party_b_address,
        "party_b_credit": raw.get("party_b_credit") or raw.get("PartyBCreditCode") or "",
        "party_b_bank": raw.get("party_b_bank") or raw.get("PartyBBank") or "",
        "party_b_account": raw.get("party_b_account") or raw.get("PartyBAccountNo") or "",
    }

def _private_customer_candidates(data):
    if isinstance(data, list):
        return data
    if isinstance(data, dict):
        return list(data.values())
    return []

def _find_private_customer(unit_name):
    if not unit_name:
        return {}
    data = read_json(getattr(config, "F_PRIVATE_CUSTOMERS", ""), [])
    for raw in _private_customer_candidates(data):
        info = _normalize_private_customer(raw)
        name = info.get("party_a_name", "")
        if name and (unit_name == name or unit_name in name or name in unit_name):
            return info
    return {}

def _create_default_template(path):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc = docx.Document()
    doc.add_heading("计算资源服务合同", 0)
    doc.add_paragraph("甲方：{{甲方单位}}")
    doc.add_paragraph("乙方：{{乙方单位}}")
    doc.add_paragraph("项目名称：{{项目名称}}")
    doc.add_paragraph("签约地点：{{签约地点}}")
    doc.add_paragraph("签约日期：{{签约日期}}")
    doc.add_heading("一、服务内容", level=1)
    doc.add_paragraph("乙方向甲方提供计算资源服务，资源清单如下：")
    doc.add_paragraph("{{资源清单}}")
    doc.add_heading("二、合同金额", level=1)
    doc.add_paragraph("合同含税总价：人民币 {{含税总价}} 元。")
    doc.add_paragraph("合同不含税总价：人民币 {{不含税总价}} 元，税率 {{税率}}。")
    doc.add_heading("三、双方信息", level=1)
    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    rows = [
        ("项目", "甲方", "乙方"),
        ("单位名称", "{{甲方单位}}", "{{乙方单位}}"),
        ("联系人", "{{甲方联系人}}", "{{乙方联系人}}"),
        ("电话", "{{甲方电话}}", "{{乙方电话}}"),
        ("地址", "{{甲方地址}}", "{{乙方地址}}"),
        ("信用代码", "{{甲方信用代码}}", "{{乙方信用代码}}"),
        ("开户行", "{{甲方开户行}}", "{{乙方开户行}}"),
        ("账号", "{{甲方账号}}", "{{乙方账号}}"),
    ]
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
    doc.add_heading("四、服务期限与付款", level=1)
    doc.add_paragraph("服务期限：{{服务期限}}。付款期限：甲方在收到乙方发票后 {{付款天数}} 日内付款。")
    doc.add_paragraph("发票类型：{{发票类型}}。")
    doc.save(path)

def _is_auto_template_without_project(path):
    try:
        doc = docx.Document(path)
        text = "\n".join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\t".join(cell.text for cell in row.cells)
        return "{{资源清单}}" in text and "{{项目名称}}" not in text
    except Exception:
        return False

def _ensure_university_contracts():
    base_names = [
        "北京大学", "清华大学", "中国人民大学", "北京航空航天大学", "北京理工大学", "中国农业大学", "北京师范大学",
        "中央民族大学", "南开大学", "天津大学", "大连理工大学", "东北大学", "吉林大学", "哈尔滨工业大学",
        "复旦大学", "同济大学", "上海交通大学", "华东师范大学", "南京大学", "东南大学", "浙江大学",
        "中国科学技术大学", "厦门大学", "山东大学", "中国海洋大学", "武汉大学", "华中科技大学", "湖南大学",
        "中南大学", "中山大学", "华南理工大学", "四川大学", "重庆大学", "电子科技大学", "西安交通大学",
        "西北工业大学", "西北农林科技大学", "兰州大学", "国防科技大学", "北京交通大学", "北京工业大学",
        "北京科技大学", "北京化工大学", "北京邮电大学", "北京林业大学", "北京中医药大学", "北京外国语大学",
        "中国传媒大学", "中央财经大学", "对外经济贸易大学", "北京体育大学", "中央音乐学院", "中国政法大学",
        "华北电力大学", "河北工业大学", "太原理工大学", "内蒙古大学", "辽宁大学", "大连海事大学",
        "延边大学", "东北师范大学", "哈尔滨工程大学", "东北农业大学", "东北林业大学", "东华大学",
        "上海财经大学", "上海大学", "苏州大学", "南京航空航天大学", "南京理工大学", "中国矿业大学",
        "河海大学", "江南大学", "南京农业大学", "中国药科大学", "南京师范大学", "安徽大学", "合肥工业大学",
        "福州大学", "南昌大学", "郑州大学", "中国地质大学", "武汉理工大学", "华中农业大学", "华中师范大学",
        "中南财经政法大学", "湖南师范大学", "暨南大学", "华南师范大学", "广西大学", "海南大学",
        "西南交通大学", "四川农业大学", "西南大学", "西南财经大学", "贵州大学", "云南大学", "西北大学",
        "西安电子科技大学", "长安大学", "陕西师范大学", "青海大学", "宁夏大学", "新疆大学", "石河子大学",
        "中国石油大学", "中国矿业大学（北京）", "中国地质大学（北京）", "中国科学院大学",
    ]
    current = read_json(config.F_UNIV_CONTRACTS, {})
    if isinstance(current, list):
        current = {}
    changed = False
    for name in base_names:
        if name not in current:
            current[name] = {
                "party_a_name": name,
                "party_a_contact": "",
                "party_a_phone": "",
                "party_a_address": "",
                "party_a_credit": "",
                "party_a_bank": "",
                "party_a_account": "",
            }
            changed = True
    if changed:
        write_json(config.F_UNIV_CONTRACTS, current)
    return current
