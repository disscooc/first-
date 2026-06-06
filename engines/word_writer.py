# word_writer.py — Word报价单（标准模板 + 云达模板）
import os, re, shutil, copy
from datetime import datetime
import docx
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn as _qn
from lxml import etree as _etree
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config
from utils import (
    safe_fn, new_path, fmt_num, fmt_money, cny, out_dir,
    price_header, comp_remark, qty_unit, stor_remark, quote_total,
    handle_err
)
from models import QuoteItem

# === 变量替换：逐run扫描替换 ===
def _sub_para(para, vars_map):
    """替换段落中的 {{var}}"""
    for run in para.runs:
        for k, v in vars_map.items():
            if k in run.text:
                run.text = run.text.replace(k, str(v) if v else "")

def _sub_cell(cell, vars_map):
    for para in cell.paragraphs:
        _sub_para(para, vars_map)

def _sub_all(doc, vars_map):
    for p in doc.paragraphs:
        _sub_para(p, vars_map)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                _sub_cell(cell, vars_map)

# === 标准模板生成 ===
def generate_standard(comp, stor, is_total, total_amt, unit, date, company, manager, phone, project):
    tpl = config.WORD_STD
    if not os.path.exists(tpl):
        raise FileNotFoundError(f"标准模板不存在: {tpl}")
    out = new_path(safe_fn(unit), "三方报价单-标准", ".docx")
    shutil.copy2(tpl, out)
    doc = docx.Document(out)

    # ========== 调试日志 ==========
    DBG = True
    def _log(msg):
        if DBG: print(f"[WORD_STD_DBG] {msg}", flush=True)
    _log("=" * 60)
    _log(f"模板: 标准模板 模式: {'total(单价)' if is_total else 'detail(数量)'}")
    _log(f"total_amt={total_amt} comp({len(comp)}) stor({len(stor)})")
    for i, it in enumerate(comp):
        _log(f"  COMP[{i}]: rtype={it.rtype} name={it.name} spec={str(it.spec)[:30]}")
        _log(f"    price={it.price} qty={it.quantity} amount={it.amount} orig_price={it.orig_price} disc={it.discount}")
    for i, it in enumerate(stor):
        _log(f"  STOR[{i}]: rtype={it.rtype} name={it.name} spec={str(it.spec)[:30]}")
        _log(f"    price={it.price} qty={it.quantity} amount={it.amount} free_quota={getattr(it,'free_quota',False)}")
    # ==============================

    # --- 段落变量 ---
    _sub_all(doc, {
        "{{报价单位}}": company,
        "{{客户单位}}": unit,
        "{{项目名称}}": project or "",
        "{{联系人}}": manager,
        "{{联系电话}}": phone,
        "{{报价日期}}": date,
    })

    tbl = doc.tables[0]
    # --- 计算资源行 (R2 是模板变量行) ---
    compute_vars = {
        "{{资源名称}}": "", "{{资源规格}}": "", "{{单价}}": "", "{{数量备注}}": "",
    }
    # 先替换模板行
    for row in tbl.rows:
        for cell in row.cells:
            for k in compute_vars:
                if k in cell.text:
                    cell.text = cell.text.replace(k, "")

    if comp:
        # 找到表头行 (第1行 index=1) 之后的数据行
        # 清空 R2 然后用它做第一行
        data_row_idx = 2  # index of {{资源名称}} row
        base_row = tbl.rows[data_row_idx]
        
        for idx, it in enumerate(comp):
            if idx == 0:
                row = base_row
            else:
                _clone_row_before(tbl, data_row_idx, data_row_idx + idx)
                row = tbl.rows[data_row_idx + idx]
            
            _set_cell(row.cells[0], it.rtype)
            _set_cell(row.cells[1], f"{it.name} {it.spec}" if it.spec else it.name)
            unit_str = qty_unit(it.rtype)
            # CPU 单价固定3位，其它资源保持原来的2位
            price_str = f"{_fmt_compute_price(it)} 元/{unit_str}"
            _set_cell(row.cells[2], price_str)
            remark = comp_remark(it.rtype)
            if not is_total and it.quantity > 0:
                # 数量：整数（无小数）
                qty_str = f"数量: {fmt_num(it.quantity, '0')} {unit_str}"
                remark = (remark + " " + qty_str).strip()
            _set_cell(row.cells[3], remark)
            # 调试日志
            _log(f"  WRITE COMP[{idx}]: row_idx={data_row_idx+idx}  col0={it.rtype}  col1={it.name}+{it.spec}  col2={price_str}  col3={remark[:40]}")
        
        # 删除多余的空行（如果模板有预留行）
        # 云达模板才有预留行，标准模板只有一个变量行

    # --- 存储资源行 ---
    # 存储行原在 R4 (index=4)，但克隆后下移到 3 + n_comp
    stor_row_idx = 3 + len(comp)
    row = tbl.rows[stor_row_idx]
    if stor:
        _set_cell(row.cells[0], stor[0].name if stor else "")
        _set_cell(row.cells[1], stor[0].spec if stor else "")
        sp = stor[0].price if stor else 0
        _set_cell(row.cells[2], f"{fmt_num(sp, '0.00')} 元/TB" if sp else "免费赠送")
    else:
        _set_cell(row.cells[0], "")
        _set_cell(row.cells[1], "")
        _set_cell(row.cells[2], "")

    # --- 总计行 ---
    total = total_amt if is_total else sum((it.amount or 0) for it in list(comp) + list(stor))
    tax_total = round(total, 2)           # total 本身就是含税总价
    no_tax = round(total / config.TAX, 2)  # 不含税价 = 含税价 / 1.06

    # 含税总计行：原 R5 → 克隆后 4 + n_comp（gridSpan 导致 cells[N] 膨胀，用 real tcs）
    tax_row = tbl.rows[4 + len(comp)]
    tax_tcs = _get_real_tcs(tax_row)
    _set_tc_text(tax_tcs[1], fmt_money(tax_total))
    # 不含税总计行：原 R6 → 克隆后 5 + n_comp
    notax_row = tbl.rows[5 + len(comp)]
    notax_tcs = _get_real_tcs(notax_row)
    _set_tc_text(notax_tcs[1], fmt_money(no_tax))

    # 确保一页输出
    _fit_to_one_page(doc)

    doc.save(out)
    return out

# === 云达模板生成 ===
def generate_yunda(comp, stor, is_total, total_amt, unit, date, company, manager, phone, project):
    tpl = config.WORD_YD
    if not os.path.exists(tpl):
        raise FileNotFoundError(f"云达模板不存在: {tpl}")
    out = new_path(safe_fn(unit), "三方报价单-云达", ".docx")
    shutil.copy2(tpl, out)
    doc = docx.Document(out)

    # ========== 调试日志 ==========
    def _log_yd(msg):
        if True: print(f"[WORD_YD_DBG] {msg}", flush=True)
    _log_yd("=" * 60)
    _log_yd(f"模板: 云达模板 模式: {'total(单价)' if is_total else 'detail(数量)'}")
    _log_yd(f"total_amt={total_amt} comp({len(comp)}) stor({len(stor)})")
    for i, it in enumerate(comp):
        _log_yd(f"  COMP[{i}]: rtype={it.rtype} name={it.name} spec={str(it.spec)[:30]}")
        _log_yd(f"    price={it.price} qty={it.quantity} amount={it.amount} orig_price={it.orig_price} disc={it.discount}")
    for i, it in enumerate(stor):
        _log_yd(f"  STOR[{i}]: rtype={it.rtype} name={it.name} spec={str(it.spec)[:30]}")
        _log_yd(f"    price={it.price} qty={it.quantity} amount={it.amount} free_quota={getattr(it,'free_quota',False)}")
    # ==============================

    # --- 段落变量 ---
    _sub_all(doc, {
        "{{报价单位}}": company,
        "{{客户单位}}": unit,
        "{{项目名称}}": project or "",
        "{{联系人}}": manager,
        "{{联系电话}}": phone,
        "{{报价日期}}": date,
    })

    tbl = doc.tables[0]
    comp_items = list(comp)
    n_comp = max(1, len(comp_items))  # 实际显示的计算资源行数（至少1行）

    # 模板固定有 2 个数据行（Row1/Row2），Row3=总计，Row4=大写
    # 注意：Row1/Row2 的 cells[0]、cells[5]、cells[6] 有跨行合并（vMerge）
    # 需要先处理行数增删
    existing_data_rows = 2

    # 如果资源数 > 2，从 Row1（模板数据行）克隆并在 Row2 前插入
    for _ in range(max(0, n_comp - existing_data_rows)):
        _clone_row_before(tbl, 1, 2)

    # 如果资源数 = 1，删除多余的第2个数据行（Row2），并把 vMerge 单元格恢复为独立单元格
    if n_comp < existing_data_rows:
        # 删除多余的空数据行
        for _ in range(existing_data_rows - n_comp):
            extra_row = tbl.rows[n_comp]  # 第 n_comp+1 行（0-indexed），即第一个多余行
            extra_row._tr.getparent().remove(extra_row._tr)
        # 处理 vMerge：删除行后，让原本 vMerge=restart 的单元格去掉 vMerge 标记
        _remove_vmerge_from_row(tbl.rows[1])

    # 计算金额
    total = total_amt if is_total else sum((it.amount or 0) for it in list(comp) + list(stor))
    tax_total = round(total, 2)
    storage_text = _storage_summary(stor)
    total_text = fmt_money(tax_total)
    cap_text = cny(tax_total)

    # --- 填写计算资源数据行 ---
    # 注意：云达模板中 cells[0]（资源类型）、cells[5]（存储）、cells[6]（总价）
    # 在原始模板中跨 Row1/Row2 垂直合并（vMerge）。
    # 为使每行资源类型独立显示，先清除所有数据行的 vMerge 标记。
    for idx in range(n_comp):
        _remove_vmerge_from_row(tbl.rows[1 + idx])

    for idx in range(n_comp):
        row = tbl.rows[1 + idx]
        tcs = _get_real_tcs(row)  # 获取本行真实的 tc 元素列表
        if idx < len(comp_items):
            it = comp_items[idx]
            # tc[0]: 资源类型（各行独立显示）
            _set_tc_text(tcs[0], it.rtype)
            # tc[1]: 配置信息（名称+规格）
            _set_tc_text(tcs[1], f"{it.name} {it.spec}" if it.spec else it.name)
            # tc[2]: 单价（CPU固定3位，其它资源保持2位）
            _set_tc_text(tcs[2], f"{_fmt_compute_price(it)}元/{qty_unit(it.rtype)}")
            # tc[3]: 数量（整数）
            _set_tc_text(tcs[3], fmt_num(it.quantity, '0') if not is_total else "")
            # tc[4]: 金额（2位小数）
            _set_tc_text(tcs[4], fmt_money(it.amount) if not is_total else "")
            # tc[5]: 存储资源（只在第一行写）
            if len(tcs) > 5:
                _set_tc_text(tcs[5], storage_text if idx == 0 else "")
            # tc[6]: 总金额（只在第一行写）
            if len(tcs) > 6:
                _set_tc_text(tcs[6], total_text if idx == 0 else "")
            # 调试日志
            _log_yd(f"  WRITE COMP[{idx}]: row={1+idx}  tc0(rtype)={it.rtype}  tc1(name+spec)={it.name}+{str(it.spec)[:20]}  tc2(price)={_fmt_compute_price(it)}  tc3(qty)={fmt_num(it.quantity, '0')}  tc4(amount)={fmt_money(it.amount)}")
        else:
            # 空行
            for tc in tcs:
                _set_tc_text(tc, "")

    # --- 填写总计行（Row 3 之后，合并单元格：cells[0,1]合并，cells[2-6]合并）---
    total_row = tbl.rows[1 + n_comp]
    total_tcs = _get_real_tcs(total_row)
    _set_tc_text(total_tcs[0], "总计")
    _set_tc_text(total_tcs[1], total_text)

    # --- 填写大写行 ---
    cap_row = tbl.rows[2 + n_comp]
    cap_tcs = _get_real_tcs(cap_row)
    _set_tc_text(cap_tcs[0], "大写")
    _set_tc_text(cap_tcs[1], cap_text)

    # --- 清除残留模板变量 ---
    _sub_all(doc, {
        "{{资源类型}}": "", "{{配置信息}}": "", "{{单价}}": "", "{{使用数量}}": "", "{{金额}}": "",
        "{{资源类型2}}": "", "{{配置信息2}}": "", "{{单价2}}": "", "{{使用数量2}}": "", "{{金额2}}": "",
        "{{存储资源}}": "", "{{总金额}}": "",
        "{{总计}}": "", "{{大写金额}}": "",
    })

    # --- 合并单元格（两种模式都做） ---
    # 1. 资源类型列：按 CPU/GPU 同类项合并
    _merge_yunda_rtype_col(tbl, comp_items, n_comp)
    # 2. 存储资源列：跨所有数据行合并
    _merge_yunda_storage_col(tbl, n_comp, is_total)
    # 3. 总金额列：仅数量模式合并（单价模式已删除该列）
    if not is_total:
        _merge_yunda_total_col(tbl, n_comp)
    # 4. 单价模式：删除使用数量、金额、总金额三列
    if is_total:
        _remove_yunda_qty_amt_cols(tbl)

    # 确保一页输出
    _fit_to_one_page(doc)

    doc.save(out)
    return out


# === 辅助函数 ===
def _set_cell(cell, text, font_name="宋体", font_size_pt=12):
    """设置 cell 文本，统一字体防止模板字体不一致。"""
    value = str(text) if text is not None else ""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if not para.runs:
        para.add_run("")
    run = para.runs[0]
    run.text = value
    # 统一字体：模板各列字体不一致，强制统一为宋体 12pt
    _set_run_font(run, font_name, font_size_pt)
    for run in para.runs[1:]:
        run.text = ""
    for extra_para in cell.paragraphs[1:]:
        for run in extra_para.runs:
            run.text = ""

def _fmt_compute_price(item):
    fmt = "0.000" if getattr(item, "rtype", "") == "CPU" else "0.00"
    return fmt_num(getattr(item, "price", 0), fmt)

def _set_run_font(run, font_name, font_size_pt):
    """设置 run 的字体（中英文统一）。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    # 设置中文字体
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    # 统一字号
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        from lxml import etree
        sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), str(font_size_pt * 2))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        from lxml import etree
        szCs = etree.SubElement(rPr, qn('w:szCs'))
    szCs.set(qn('w:val'), str(font_size_pt * 2))
    # 清除加粗（数据行不应加粗）
    b = rPr.find(qn('w:b'))
    if b is not None:
        rPr.remove(b)
    bCs = rPr.find(qn('w:bCs'))
    if bCs is not None:
        rPr.remove(bCs)

def _get_real_tcs(row):
    """获取 row 中真实的 <w:tc> 元素列表（不受合并影响）。"""
    from docx.oxml.ns import qn
    return row._tr.findall(qn('w:tc'))

def _set_tc_text(tc, text, font_name="宋体", font_size_pt=12):
    """直接在 <w:tc> XML 元素上设置文本，并统一字体。"""
    from docx.oxml.ns import qn
    from lxml import etree
    value = str(text) if text is not None else ""
    # 找到第一个段落中的第一个 run
    paras = tc.findall('.//' + qn('w:p'))
    if not paras:
        return
    para = paras[0]
    runs = para.findall(qn('w:r'))
    if runs:
        # 找到第一个 run，修改其 w:t
        r = runs[0]
        wt = r.find(qn('w:t'))
        if wt is None:
            wt = etree.SubElement(r, qn('w:t'))
        wt.text = value
        if value and (value[0] == ' ' or value[-1] == ' '):
            wt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        # 统一字体：强制宋体 12pt
        _set_xml_run_font(r, font_name, font_size_pt)
        # 清空其余 run
        for extra_r in runs[1:]:
            wt2 = extra_r.find(qn('w:t'))
            if wt2 is not None:
                wt2.text = ""
    else:
        # 没有 run，创建一个
        r = etree.SubElement(para, qn('w:r'))
        wt = etree.SubElement(r, qn('w:t'))
        wt.text = value
        _set_xml_run_font(r, font_name, font_size_pt)
    # 清空其余段落的内容
    for extra_para in paras[1:]:
        for r in extra_para.findall(qn('w:r')):
            wt = r.find(qn('w:t'))
            if wt is not None:
                wt.text = ""

def _set_xml_run_font(r, font_name, font_size_pt):
    """在 XML run 元素上设置统一字体（中英文同字体同字号）。"""
    from docx.oxml.ns import qn
    from lxml import etree
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(r, qn('w:rPr'))
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    sz_val = str(font_size_pt * 2)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), sz_val)
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = etree.SubElement(rPr, qn('w:szCs'))
    szCs.set(qn('w:val'), sz_val)
    # 清除加粗
    for tag in ('w:b', 'w:bCs'):
        el = rPr.find(qn(tag))
        if el is not None:
            rPr.remove(el)

def _remove_vmerge_from_row(row):
    """移除某行中所有 vMerge 标记（用于删除合并行后恢复独立行）。"""
    from docx.oxml.ns import qn
    for tc in _get_real_tcs(row):
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            vm = tcPr.find(qn('w:vMerge'))
            if vm is not None:
                tcPr.remove(vm)

def _clone_row_before(tbl, template_idx, before_idx):
    new_tr = copy.deepcopy(tbl.rows[template_idx]._tr)
    tbl.rows[before_idx]._tr.addprevious(new_tr)


def _storage_summary(stor):
    parts = []
    for item in stor:
        if item.free_quota:
            # 免费赠送：显示固定赠送量 FREE_STORAGE（0.5T），不管 quantity 是多少
            free = fmt_num(config.FREE_STORAGE, '0.##')
            parts.append(f"{item.name}：免费赠送{free}T，\n免费额度内不收费")
        else:
            # 付费存储：显示实际容量（整数）
            qty = fmt_num(item.quantity, "0") if item.quantity else ""
            parts.append(f"{item.name} {qty}T")
    return "\n".join(parts)

def copy_cell_style(src_cell, dst_cell):
    """复制单元格样式"""
    # python-docx 不直接支持复制cell，用简单方式
    for pi, sp in enumerate(src_cell.paragraphs):
        if pi < len(dst_cell.paragraphs):
            dp = dst_cell.paragraphs[pi]
        else:
            dp = dst_cell.add_paragraph()
        dp.alignment = sp.alignment
        for ri, sr in enumerate(sp.runs):
            if ri < len(dp.runs):
                dr = dp.runs[ri]
            else:
                dr = dp.add_run("")
            dr.bold = sr.bold
            dr.italic = sr.italic
            dr.underline = sr.underline
            dr.font.size = sr.font.size
            dr.font.name = sr.font.name


def _remove_yunda_qty_amt_cols(tbl):
    """单价模式下删除云达模板的"使用数量"(col3)、"金额"(col4)、"总金额"(col6)三列。
    同时更新 tblGrid、表头/数据行的 tc、以及总计/大写行的 gridSpan。
    """
    from docx.oxml.ns import qn
    COLS_TO_DEL = {3, 4, 6}

    # 1. 删除 tblGrid 中的 gridCol
    tblGrid = tbl._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        gridCols = tblGrid.findall(qn('w:gridCol'))
        for ci in sorted(COLS_TO_DEL, reverse=True):
            if ci < len(gridCols):
                tblGrid.remove(gridCols[ci])

    # 2. 遍历每一行处理 tc
    for row in tbl.rows:
        tcs = row._tr.findall(qn('w:tc'))
        if len(tcs) >= 7:
            # 普通行（原 7 列）：删除 tc[3]、tc[4]、tc[6]（从后往前删避免索引偏移）
            for ci in sorted(COLS_TO_DEL, reverse=True):
                if ci < len(tcs):
                    tcs[ci].getparent().remove(tcs[ci])
        else:
            # 合并行（总计/大写，2 个 tc）：调整 gridSpan
            # 原布局：tc[0] gridSpan=2 覆盖列 0-1 / tc[1] gridSpan=5 覆盖列 2-6
            # 删除列 3,4,6 后：tc[0] 仍 2（未变）/ tc[1] 5-3=2
            for i, tc in enumerate(tcs):
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    gs = tcPr.find(qn('w:gridSpan'))
                    if gs is not None:
                        if i == 0:
                            # tc[0] 覆盖列 0-1，无被删列 → 保持 2
                            gs.set(qn('w:val'), '2')
                        else:
                            # tc[1] 覆盖列 2-6，删除列 3,4,6 → 剩 2
                            gs.set(qn('w:val'), '2')


# 临时文件：新的合并函数，将被插入到 word_writer.py
# 用法：由 apply_patch.py 读取并替换

def _merge_yunda_rtype_col(tbl, comp_items, n_comp):
    """资源类型列（col 0）按 CPU/GPU 同类项垂直合并（同时支持单价/数量模式）。"""
    from docx.oxml.ns import qn
    from lxml import etree

    # 按 rtype 分组连续行
    groups = []
    current_group = []
    for idx in range(n_comp):
        rtype = comp_items[idx].rtype if idx < len(comp_items) else None
        if not current_group or current_group[-1][1] == rtype:
            current_group.append((idx, rtype))
        else:
            groups.append(current_group)
            current_group = [(idx, rtype)]
    if current_group:
        groups.append(current_group)

    for group in groups:
        if len(group) <= 1:
            continue  # 单行无需合并
        for gi, (row_idx, _rtype) in enumerate(group):
            row = tbl.rows[1 + row_idx]
            tcs = _get_real_tcs(row)
            if not tcs:
                continue
            tc = tcs[0]  # 资源类型 = col 0（两种模式都是 col 0）
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = etree.SubElement(tc, qn('w:tcPr'))
                tc.insert(0, tcPr)
            # 清除旧的 vMerge
            old_vm = tcPr.find(qn('w:vMerge'))
            if old_vm is not None:
                tcPr.remove(old_vm)
            vm = etree.SubElement(tcPr, qn('w:vMerge'))
            vm.set(qn('w:val'), 'restart' if gi == 0 else 'continue')


def _merge_yunda_storage_col(tbl, n_comp, is_total):
    """存储资源列跨所有数据行垂直合并，不留空白格。
    is_total=True（单价模式）：列已删除重组，存储列 = tcs[3]
    is_total=False（数量模式）：7列完整，存储列 = tcs[5]
    """
    from docx.oxml.ns import qn
    from lxml import etree

    for idx in range(n_comp):
        row = tbl.rows[1 + idx]
        tcs = _get_real_tcs(row)
        if not tcs:
            continue
        # 根据模式选择列索引
        if is_total:
            col_idx = 3  # 单价模式：删除了 col 3,4,6，原 col 5 变为 col 3
        else:
            col_idx = 5  # 数量模式：7列完整，存储 = col 5
        if len(tcs) <= col_idx:
            continue
        tc = tcs[col_idx]
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            tc.insert(0, tcPr)
        # 清除旧的 vMerge
        old_vm = tcPr.find(qn('w:vMerge'))
        if old_vm is not None:
            tcPr.remove(old_vm)
        vm = etree.SubElement(tcPr, qn('w:vMerge'))
        vm.set(qn('w:val'), 'restart' if idx == 0 else 'continue')


def _merge_yunda_total_col(tbl, n_comp):
    """数量模式下：总金额列（col 6）跨所有数据行垂直合并，只在第一行显示。"""
    from docx.oxml.ns import qn
    from lxml import etree

    for idx in range(n_comp):
        row = tbl.rows[1 + idx]
        tcs = _get_real_tcs(row)
        if not tcs or len(tcs) <= 6:
            continue
        tc = tcs[6]  # 总金额 = col 6（仅数量模式存在）
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            tc.insert(0, tcPr)
        # 清除旧的 vMerge
        old_vm = tcPr.find(qn('w:vMerge'))
        if old_vm is not None:
            tcPr.remove(old_vm)
        vm = etree.SubElement(tcPr, qn('w:vMerge'))
        vm.set(qn('w:val'), 'restart' if idx == 0 else 'continue')


# === 页面适配：确保一页输出 ===

def _fit_to_one_page(doc):
    """压缩行高和段落间距，确保报价单不超一页。"""
    section = doc.sections[0]
    usable_emu = int((section.page_height - section.top_margin - section.bottom_margin) * 0.95)

    # 1. 确保首标题加粗（保留模板字号），其余段落不压缩字距
    for p in doc.paragraphs:
        if p.text.strip():
            for run in p.runs:
                run.bold = True
            break  # 只处理第一个非空段落（标题）

    # 2. 估算当前总高度，仅在超页时压缩
    tbl = doc.tables[0]
    total_emu = _estimate_doc_height(doc, tbl)
    if total_emu <= usable_emu:
        return  # 已在一页内，不做改动

    # 3. 超页：按比例缩放所有行高，不低于 12pt（240 twips）
    scale = usable_emu / total_emu
    # 读取每行当前高度，缩放后取 max(240, scaled)
    row_heights = []
    for row in tbl.rows:
        h = _get_tr_height(row)
        row_heights.append(h)

    # 统一缩放
    for i, row in enumerate(tbl.rows):
        new_h = max(240, int(row_heights[i] * scale))
        _set_tr_height(row, new_h, 'exact')


def _set_tr_height(row, val_twips, rule='atLeast'):
    """设置表格行高（val_twips 单位为 1/20 pt）。"""
    trPr = row._tr.find(_qn('w:trPr'))
    if trPr is None:
        trPr = _etree.SubElement(row._tr, _qn('w:trPr'))
        row._tr.insert(0, trPr)
    trHeight = trPr.find(_qn('w:trHeight'))
    if trHeight is None:
        trHeight = _etree.SubElement(trPr, _qn('w:trHeight'))
    trHeight.set(_qn('w:val'), str(val_twips))
    if rule == 'exact':
        trHeight.set(_qn('w:hRule'), 'exact')
    else:
        try:
            del trHeight.attrib[_qn('w:hRule')]
        except KeyError:
            pass


def _get_tr_height(row):
    """读取表格行高（twips），默认返回 400（20pt）。"""
    trPr = row._tr.find(_qn('w:trPr'))
    if trPr is not None:
        th = trPr.find(_qn('w:trHeight'))
        if th is not None:
            return int(th.get(_qn('w:val'), '400'))
    return 400


def _estimate_doc_height(doc, tbl):
    """估算文档总高度（EMU）。"""
    total = 0
    # 非空段落：每个约 0.6cm（216000 EMU）
    for p in doc.paragraphs:
        if p.text.strip():
            total += 216000
    # 表格行高度
    for row in tbl.rows:
        trPr = row._tr.find(_qn('w:trPr'))
        if trPr is not None:
            th = trPr.find(_qn('w:trHeight'))
            if th is not None:
                total += int(th.get(_qn('w:val'), '300')) * 635
            else:
                total += 300 * 635
        else:
            total += 300 * 635
    return total
