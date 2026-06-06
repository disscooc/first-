# engines/contract_scanner.py -- 扫描历史合同并辅助补充库
import json
import re
import shutil
from dataclasses import dataclass, asdict, field
from datetime import datetime
from pathlib import Path

import docx

import config
from utils import read_json, write_json
from engines.library_manager import list_party_b_names


CONTRACT_KEYWORDS = ("合同", "协议", "服务合同", "采购合同", "机时服务合同", "算力服务合同", "技术服务合同")
SKIP_KEYWORDS = ("报价单", "机时服务清单", "清单", "quote", "quotation", "price", "报价")

A_ROLES = ("甲方", "甲方单位", "用户单位", "采购方", "委托方", "买方")
B_ROLES = ("乙方", "乙方单位", "报价单位", "服务方", "受托方", "卖方")
A_CONTEXT = A_ROLES + ("甲方联系人", "用户", "买方联系人", "委托方联系人")
B_CONTEXT = B_ROLES + ("乙方联系人", "服务方联系人", "卖方联系人", "受托方联系人")

UNIT_KEYWORDS = ("大学", "学院", "研究院", "研究所", "公司", "有限公司", "中心", "医院", "实验室", "集团", "局", "所", "厂", "学校")
UNIT_SUFFIX = "|".join(sorted(UNIT_KEYWORDS + ("事务所", "委员会", "办公室"), key=len, reverse=True))
PHONE_RE = re.compile(
    r"(?:(?:1[3-9]\d{9})|(?:0\d{2,5}[-\s]?\d{6,8}(?:[-\s]?\d{1,6})?)|(?:\+?852[-\s]?\d{8})|(?:00852[-\s]?\d{8}))"
)
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
CREDIT_RE = re.compile(r"(?<![0-9A-Z])([0-9A-Z]{18})(?![0-9A-Z])")
ACCOUNT_RE = re.compile(r"^[0-9][0-9\-\s]{7,39}$")

BODY_WORDS = ("双方", "合同", "服务期", "重叠", "甲方", "乙方", "如果", "则", "本合同", "条款", "约定")
BAD_UNIT_WORDS = (
    "协议", "本协议", "合同", "本合同", "条款", "规定", "工作", "服务期", "签署", "双方", "重叠", "期间",
    "除本", "如双方", "履行", "支付", "费用", "金额", "税率", "开票", "发票", "盖章", "签字", "日期",
)
BAD_CONTACT_WORDS = ("签字", "盖章", "日期", "代表", "法定代表人", "授权代表", "经办人签字", "负责人签字")
FIELD_LABELS = (
    "甲方", "乙方", "甲方单位", "乙方单位", "用户单位", "报价单位", "采购方", "委托方", "服务方", "受托方",
    "单位名称", "注册地址", "住所", "地址", "单位地址", "统一社会信用代码", "信用代码", "税号",
    "开户行", "开户银行", "账号", "银行账号", "开户账号", "联系人", "联系人姓名", "甲方联系人", "乙方联系人",
    "经办人", "业务联系人", "电话", "联系电话", "手机", "联系方式", "邮箱", "电子邮箱", "通讯地址", "联系地址", "邮寄地址",
)


@dataclass
class ContractLine:
    index: int
    text: str
    cells: list[str] = field(default_factory=list)
    source: str = ""


@dataclass
class ContractScanResult:
    path: str
    file_name: str
    party_a_name: str = ""
    party_a_address: str = ""
    party_a_credit: str = ""
    party_a_bank: str = ""
    party_a_account: str = ""
    party_a_contact: str = ""
    party_a_phone: str = ""
    party_a_email: str = ""
    party_a_mailing_address: str = ""
    party_b_name: str = ""
    party_b_address: str = ""
    party_b_credit: str = ""
    party_b_bank: str = ""
    party_b_account: str = ""
    party_b_contact: str = ""
    party_b_phone: str = ""
    party_b_email: str = ""
    party_b_mailing_address: str = ""
    pending_info: str = ""
    confidence: int = 0
    recognition_status: str = "待确认"
    operation: str = "导入"
    debug_notes: list[str] = field(default_factory=list)


def should_scan_contract(path):
    name = path.name.lower()
    if path.suffix.lower() != ".docx":
        return False, "非docx合同"
    if path.name.startswith("~$"):
        return False, "临时文件"
    if any(keyword.lower() in name for keyword in SKIP_KEYWORDS):
        return False, "报价/清单文件"
    if not any(keyword.lower() in name for keyword in CONTRACT_KEYWORDS):
        return False, "文件名不像合同"
    return True, ""


def scan_contract_folder(folder):
    folder = Path(folder)
    results = []
    skipped = []
    unrecognized = []
    for path in folder.rglob("*"):
        if not path.is_file():
            continue
        if path.suffix.lower() not in (".docx", ".doc"):
            skipped.append((str(path), "非合同Word文件"))
            continue
        can_scan, reason = should_scan_contract(path)
        if not can_scan:
            skipped.append((str(path), reason))
            continue
        try:
            result = scan_contract_docx(path)
        except Exception as exc:
            unrecognized.append((str(path), f"读取失败：{exc}"))
            continue
        if _has_any_core_info(result):
            results.append(result)
        else:
            unrecognized.append((str(path), "未识别到甲乙方信息"))
    low_confidence = sum(1 for item in results if item.confidence < 40)
    pending_confidence = sum(1 for item in results if 40 <= item.confidence < 70)
    return {
        "results": results,
        "scanned_count": len(results),
        "success_count": len([item for item in results if item.confidence >= 70]),
        "pending_count": pending_confidence,
        "low_confidence_count": low_confidence,
        "skipped_count": len(skipped),
        "unrecognized_count": len(unrecognized),
        "skipped": skipped,
        "unrecognized": unrecognized,
    }


def scan_contract_docx(path):
    lines = _docx_lines(path)
    side_by_index = _side_context(lines)
    result = ContractScanResult(path=str(path), file_name=Path(path).name)
    pending = []

    result.party_a_name = _extract_party_name(lines, "party_a", side_by_index, result.debug_notes)
    result.party_b_name = _extract_party_name(lines, "party_b", side_by_index, result.debug_notes)

    _fill_company_fields(result, "party_a", lines, side_by_index)
    _fill_company_fields(result, "party_b", lines, side_by_index)
    _fill_contact_fields(result, "party_a", lines, side_by_index)
    _fill_contact_fields(result, "party_b", lines, side_by_index)
    _assign_emails(lines, result, pending)
    _collect_unconfirmed_phones(lines, result, pending)

    _finalize_result(result, pending)
    return result


def _docx_lines(path):
    document = docx.Document(path)
    rows = []
    index = 0
    for para in document.paragraphs:
        text = _clean(para.text)
        if text:
            rows.append(ContractLine(index=index, text=text, cells=[text], source="paragraph"))
            index += 1
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            cells = [_clean(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if not cells:
                continue
            rows.append(
                ContractLine(
                    index=index,
                    text="\t".join(cells),
                    cells=cells,
                    source=f"table{table_index + 1}-row{row_index + 1}",
                )
            )
            index += 1
    return rows


def _clean(value):
    return re.sub(r"\s+", " ", str(value or "")).strip(" ：:，,。；;")


def _compact(value):
    return re.sub(r"\s+", "", str(value or ""))


def _label_pattern(label):
    return r"\s*".join(re.escape(char) for char in label)


def _contains_label(text, label):
    return re.search(_label_pattern(label), text or "") is not None


def _contains_any_label(text, labels):
    return any(_contains_label(text, label) for label in labels)


def _side_context(lines):
    side_by_index = {}
    current = None
    for line in lines:
        explicit = _explicit_side(line.text)
        if explicit:
            current = explicit
        side_by_index[line.index] = current
    return side_by_index


def _explicit_side(text):
    compact = _compact(text)
    has_a = any(_compact(word) in compact for word in A_CONTEXT)
    has_b = any(_compact(word) in compact for word in B_CONTEXT)
    if has_a and not has_b:
        return "party_a"
    if has_b and not has_a:
        return "party_b"
    return None


def _applies_to_side(line, side, side_by_index):
    explicit = _explicit_side(line.text)
    if explicit:
        return explicit == side
    return side_by_index.get(line.index) == side


def _is_body_clause(text):
    text = _clean(text)
    if re.match(r"^\d+(\.\d+)*[\u4e00-\u9fa5]", text) and any(word in text for word in BAD_UNIT_WORDS):
        return True
    if re.match(r"^\d+(\.\d+)*[、.．]", text) and any(word in text for word in BODY_WORDS):
        return True
    if re.match(r"^第?[一二三四五六七八九十百]+[章节条款]", text):
        return True
    return len(text) > 70 and any(word in text for word in BODY_WORDS + BAD_UNIT_WORDS)


def _extract_party_name(lines, side, side_by_index, debug):
    labels = A_ROLES if side == "party_a" else B_ROLES
    field_label = "甲方单位" if side == "party_a" else "乙方单位"
    for line in lines:
        if _is_body_clause(line.text):
            candidates = [line.text]
            candidates.extend(candidate for candidate in _unit_candidates(line.text) if candidate != line.text)
            for candidate in candidates:
                debug.append(_debug(field_label, candidate, "正文条款", line, "丢弃：疑似合同正文，不从正文猜测单位"))
            continue
        if not _applies_to_side(line, side, side_by_index):
            continue
        if "联系人" in line.text:
            continue
        matched = _matched_label(line.text, labels)
        value = _value_from_line(line, labels)
        if value is None:
            for candidate in _unit_candidates(line.text):
                debug.append(_debug(field_label, candidate, "非明确甲乙方标签", line, "丢弃：单位必须来自明确标签或表格字段"))
            continue
        ok, unit, reason = _valid_unit_name(value, explicit=True)
        if ok:
            debug.append(_debug(field_label, unit, matched or labels[0], line))
            return unit
        for candidate in _unit_candidates(value) or [value]:
            debug.append(_debug(field_label, candidate, matched or labels[0], line, f"丢弃：{reason}"))
    return ""


def _fill_company_fields(result, prefix, lines, side_by_index):
    debug = result.debug_notes
    field_specs = [
        ("address", ("注册地址", "住所", "单位地址", "地址"), _valid_address),
        ("credit", ("统一社会信用代码", "信用代码", "税号"), _valid_credit),
        ("bank", ("开户行", "开户银行"), _valid_bank),
        ("account", ("账号", "银行账号", "开户账号"), _valid_account),
    ]
    for attr, labels, validator in field_specs:
        value = _extract_labeled_value(lines, prefix, side_by_index, labels, validator, debug, attr)
        setattr(result, f"{prefix}_{attr}", value)


def _fill_contact_fields(result, prefix, lines, side_by_index):
    debug = result.debug_notes
    has_unit = bool(getattr(result, f"{prefix}_name", ""))
    contact = _extract_labeled_value(
        lines,
        prefix,
        side_by_index,
        ("甲方联系人", "乙方联系人", "联系人姓名", "联系人", "经办人", "业务联系人"),
        _valid_contact,
        debug,
        "contact",
        strict_side=True,
        allow_section_side=has_unit,
    )
    phone = _extract_labeled_value(
        lines,
        prefix,
        side_by_index,
        ("联系电话", "电话", "手机", "联系方式"),
        _valid_phone,
        debug,
        "phone",
        strict_side=True,
        explicit_side_only=not has_unit,
        allow_section_side=has_unit,
    )
    mailing = _extract_labeled_value(
        lines,
        prefix,
        side_by_index,
        ("通讯地址", "联系地址", "邮寄地址"),
        _valid_address,
        debug,
        "mailing_address",
        strict_side=True,
        allow_section_side=has_unit,
    )
    setattr(result, f"{prefix}_contact", contact)
    setattr(result, f"{prefix}_phone", phone)
    setattr(result, f"{prefix}_mailing_address", mailing)


def _extract_labeled_value(
    lines,
    side,
    side_by_index,
    labels,
    validator,
    debug,
    field_name,
    strict_side=False,
    explicit_side_only=False,
    allow_section_side=False,
):
    for position, line in enumerate(lines):
        if _is_body_clause(line.text):
            if _contains_any_label(line.text, labels):
                debug.append(_debug(field_name, line.text, "正文条款", line, "丢弃：正文条款"))
            continue
        if explicit_side_only:
            line_side = _explicit_side(line.text)
            if line_side != side:
                if _contains_any_label(line.text, labels):
                    debug.append(_debug(field_name, line.text, "归属不明确", line, "丢弃：行内未明确甲方或乙方"))
                continue
        elif strict_side:
            line_side = _strict_side_for_line(lines, position, line)
            if not line_side and allow_section_side:
                line_side = side_by_index.get(line.index)
            if line_side != side:
                if _contains_any_label(line.text, labels):
                    debug.append(_debug(field_name, line.text, "归属不明确", line, "丢弃：无法明确属于甲方或乙方"))
                continue
        elif not _applies_to_side(line, side, side_by_index):
            continue
        if field_name == "phone" and _contains_any_label(line.text, ("注册地址", "单位地址", "住所", "地址", "统一社会信用代码", "信用代码", "税号", "账号", "银行账号", "开户账号")):
            debug.append(_debug(field_name, line.text, "非联系人电话", line, "丢弃：注册地址/账号/信用代码行不作为联系人电话"))
            continue
        value = _value_from_line(line, labels)
        if value is None:
            continue
        value = _trim_field_value(value)
        ok, normalized, reason = validator(value)
        if ok:
            debug.append(_debug(field_name, normalized, _matched_label(line.text, labels), line))
            return normalized
        debug.append(_debug(field_name, value, _matched_label(line.text, labels), line, f"丢弃：{reason}"))
    return ""


def _strict_side_for_line(lines, position, line):
    explicit = _explicit_side(line.text)
    if explicit:
        return explicit
    context_lines = lines[max(0, position - 3): position + 4]
    context = "\n".join(item.text for item in context_lines)
    return _side_from_context(context)


def _value_from_line(line, labels):
    for label in labels:
        for index, cell in enumerate(line.cells):
            if not _contains_label(cell, label):
                continue
            after = _after_label(cell, label)
            if after:
                return after
            if index + 1 < len(line.cells):
                return line.cells[index + 1]
        if _contains_label(line.text, label):
            after = _after_label(line.text, label)
            if after:
                return after
    return None


def _after_label(text, label):
    pattern = rf"{_label_pattern(label)}\s*(?:[、,，]\s*{_label_pattern('电话')})?\s*[:：]?\s*(.*)"
    match = re.search(pattern, text)
    if not match:
        return ""
    return _clean(match.group(1))


def _trim_field_value(value):
    value = _clean(value)
    cut_pos = None
    for label in FIELD_LABELS:
        match = re.search(_label_pattern(label), value)
        if match and match.start() > 0:
            cut_pos = match.start() if cut_pos is None else min(cut_pos, match.start())
    if cut_pos is not None:
        value = value[:cut_pos]
    value = re.split(r"[\t；;。]", value)[0]
    return _clean(value)


def _unit_candidate(text):
    if not text:
        return ""
    for item in _unit_candidates(text):
        ok, normalized, _reason = _valid_unit_name(item, explicit=False)
        if ok:
            return normalized
    return ""


def _unit_candidates(text):
    if not text:
        return []
    pattern = rf"[\u4e00-\u9fa5A-Za-z0-9（）()·\-]{{1,60}}?(?:{UNIT_SUFFIX})"
    return [_clean(item) for item in re.findall(pattern, text) if _clean(item)]


def _valid_unit_name(value, explicit=False):
    value = _trim_field_value(value)
    if not value:
        return False, "", "为空"
    if _is_body_clause(value):
        return False, "", "疑似合同正文"
    if any(word in value for word in BAD_UNIT_WORDS):
        hits = "/".join(word for word in BAD_UNIT_WORDS if word in value)
        return False, "", f"包含正文关键词：{hits}"
    if PHONE_RE.search(value) or EMAIL_RE.search(value):
        return False, "", "包含电话或邮箱"
    if re.match(r"^\d", value):
        return False, "", "以数字开头，疑似条款编号"
    if len(value) > 80:
        return False, "", "单位名称过长"
    candidates = _unit_candidates(value)
    if explicit and any(word in value for word in UNIT_KEYWORDS):
        pass
    elif candidates:
        value = max(candidates, key=len)
    elif not explicit:
        return False, "", "不含单位关键词"
    if not any(word in value for word in UNIT_KEYWORDS):
        return False, "", "不含单位关键词"
    return True, _clean(value), ""


def _looks_like_unit(value):
    ok, _normalized, _reason = _valid_unit_name(value, explicit=False)
    return ok


def _valid_address(value):
    value = _clean(value)
    if not value:
        return False, "", "为空"
    if _is_body_clause(value):
        return False, "", "疑似合同正文"
    if len(value) > 120:
        return False, "", "过长"
    return True, value, ""


def _valid_credit(value):
    match = CREDIT_RE.search(value or "")
    if not match:
        return False, "", "不是18位统一社会信用代码"
    return True, match.group(0), ""


def _valid_bank(value):
    value = _clean(value)
    if not value:
        return False, "", "为空"
    if _is_body_clause(value):
        return False, "", "疑似合同正文"
    if len(value) > 80:
        return False, "", "过长"
    return True, value, ""


def _valid_account(value):
    value = _clean(value)
    if any(word in value for word in BODY_WORDS):
        return False, "", "包含正文关键词"
    if len(re.findall(r"[\u4e00-\u9fa5]", value)) > 2:
        return False, "", "包含大量中文"
    compact = re.sub(r"[\s\-]", "", value)
    if not compact.isdigit() or len(compact) <= 8:
        return False, "", "不像银行账号"
    if not ACCOUNT_RE.match(value):
        return False, "", "账号格式不合法"
    return True, value, ""


def _valid_contact(value):
    value = _clean(value)
    value = re.split(r"(?:联系电话|电话|手机|联系方式|邮箱|电子邮箱|通讯地址|联系地址|邮寄地址)[:：]?", value)[0]
    value = _clean(value)
    if not value:
        return False, "", "为空"
    if any(word in value for word in BAD_CONTACT_WORDS):
        hits = "/".join(word for word in BAD_CONTACT_WORDS if word in value)
        return False, "", f"联系人禁用词：{hits}"
    if not re.fullmatch(r"[\u4e00-\u9fa5]{2,4}", value):
        return False, "", "联系人必须是2-4个中文字符"
    if PHONE_RE.search(value) or EMAIL_RE.search(value) or _looks_like_unit(value):
        return False, "", "不是联系人姓名"
    return True, value, ""


def _valid_phone(value):
    match = PHONE_RE.search(value or "")
    if not match:
        return False, "", "未找到电话"
    return True, _clean(match.group(0)), ""


def _assign_emails(lines, result, pending):
    found = set()
    for index, line in enumerate(lines):
        for match in EMAIL_RE.finditer(line.text):
            email = match.group(0)
            if email in found:
                continue
            found.add(email)
            context_lines = lines[max(0, index - 3): index + 4]
            context = "\n".join(item.text for item in context_lines)
            side = _side_from_context(context)
            if not side:
                side = _side_from_email_domain(email, result)
            if side == "party_a" and not result.party_a_email:
                result.party_a_email = email
                result.debug_notes.append(_debug("party_a_email", email, "邮箱上下文", line, "写入甲方邮箱"))
            elif side == "party_b" and not result.party_b_email:
                result.party_b_email = email
                result.debug_notes.append(_debug("party_b_email", email, "邮箱上下文", line, "写入乙方邮箱"))
            else:
                pending.append(f"发现邮箱：{email}")
                result.debug_notes.append(_debug("pending_email", email, "邮箱上下文不明确", line, "放入待确认"))


def _side_from_context(context):
    has_a = any(word in context for word in A_CONTEXT)
    has_b = any(word in context for word in B_CONTEXT)
    if has_a and not has_b:
        return "party_a"
    if has_b and not has_a:
        return "party_b"
    return None


def _side_from_email_domain(email, result):
    domain = email.split("@", 1)[-1].lower()
    a_name = result.party_a_name or ""
    if domain.endswith(".edu.cn") and any(word in a_name for word in ("大学", "学院", "研究院", "研究所")):
        return "party_a"
    if "cup.edu.cn" in domain and "石油" in a_name:
        return "party_a"
    return None


def _collect_unconfirmed_phones(lines, result, pending):
    known = {result.party_a_phone, result.party_b_phone}
    for line in lines:
        if _contains_any_label(line.text, ("注册地址", "单位地址", "住所", "地址", "统一社会信用代码", "信用代码", "税号", "账号", "银行账号", "开户账号")):
            continue
        for match in PHONE_RE.finditer(line.text):
            phone = _clean(match.group(0))
            if phone and phone not in known:
                pending.append(f"发现电话：{phone}")
                result.debug_notes.append(_debug("pending_phone", phone, "电话归属不明确", line, "放入待确认"))
                known.add(phone)


def _matched_label(text, labels):
    for label in labels:
        if _contains_label(text, label):
            return label
    return ""


def _debug(field_name, value, keyword, line, action="写入"):
    return (
        f"字段：{field_name}\n"
        f"候选内容：{value}\n"
        f"来源关键词：{keyword}\n"
        f"上下文：{line.text}\n"
        f"判断：{action}\n"
    )


def _has_any_core_info(result):
    return bool(
        result.party_a_name
        or result.party_b_name
        or result.party_a_contact
        or result.party_b_contact
        or result.pending_info
        or result.debug_notes
    )


def _confidence(result):
    score = 0
    for field_name, points in (
        ("party_a_name", 18), ("party_b_name", 18),
        ("party_a_credit", 10), ("party_b_credit", 10),
        ("party_a_contact", 8), ("party_b_contact", 8),
        ("party_a_phone", 6), ("party_b_phone", 6),
        ("party_a_email", 4), ("party_b_email", 4),
        ("party_a_bank", 5), ("party_b_bank", 5),
        ("party_a_account", 3), ("party_b_account", 3),
    ):
        if getattr(result, field_name, ""):
            score += points
    return min(score, 100)


SCAN_FIELD_NAMES = (
    "party_a_name", "party_a_address", "party_a_credit", "party_a_bank", "party_a_account",
    "party_a_contact", "party_a_phone", "party_a_email", "party_a_mailing_address",
    "party_b_name", "party_b_address", "party_b_credit", "party_b_bank", "party_b_account",
    "party_b_contact", "party_b_phone", "party_b_email", "party_b_mailing_address",
)


def _finalize_result(result, pending):
    confidence = _confidence(result)
    result.confidence = confidence
    if confidence < 40:
        moved = []
        for field_name in SCAN_FIELD_NAMES:
            value = getattr(result, field_name, "")
            if value:
                moved.append(f"{_field_display_name(field_name)}：{value}")
                setattr(result, field_name, "")
        if moved:
            pending.append("低置信度候选，未自动填入：\n" + "\n".join(moved))
            result.debug_notes.append(
                "低置信度处理：识别置信度低于40，已清空自动字段，仅保留到未确认信息。\n"
                + "\n".join(moved)
                + "\n"
            )
        result.recognition_status = "低置信度"
        result.operation = "跳过"
    elif confidence < 70:
        result.recognition_status = "待确认"
        result.operation = "待确认"
    else:
        result.recognition_status = "高置信度"
        result.operation = "导入"
    result.pending_info = "\n".join(dict.fromkeys(pending))


def _field_display_name(field_name):
    names = {
        "party_a_name": "甲方单位",
        "party_a_address": "甲方注册地址",
        "party_a_credit": "甲方信用代码",
        "party_a_bank": "甲方开户行",
        "party_a_account": "甲方账号",
        "party_a_contact": "甲方联系人",
        "party_a_phone": "甲方电话",
        "party_a_email": "甲方邮箱",
        "party_a_mailing_address": "甲方通讯地址",
        "party_b_name": "乙方单位",
        "party_b_address": "乙方注册地址",
        "party_b_credit": "乙方信用代码",
        "party_b_bank": "乙方开户行",
        "party_b_account": "乙方账号",
        "party_b_contact": "乙方联系人",
        "party_b_phone": "乙方电话",
        "party_b_email": "乙方邮箱",
        "party_b_mailing_address": "乙方通讯地址",
    }
    return names.get(field_name, field_name)


def backup_libraries():
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_dir = Path(config.USER_DATA_DIR) / f"library_backup_{stamp}"
    backup_dir.mkdir(parents=True, exist_ok=True)
    for path in (
        config.F_PRIVATE_CUSTOMERS,
        config.F_PRIVATE_CONTACTS,
        config.F_PARTY_B_LIBRARY,
        getattr(config, "F_SELLER_CONTACTS", ""),
    ):
        if path and Path(path).exists():
            shutil.copy2(path, backup_dir / Path(path).name)
    return str(backup_dir)


def import_scan_results(results):
    backup_dir = backup_libraries()
    stats = {
        "imported_units": 0,
        "imported_contacts": 0,
        "duplicates": 0,
        "backup_dir": backup_dir,
    }
    for raw in results:
        item = raw if isinstance(raw, ContractScanResult) else ContractScanResult(**raw)
        if item.operation != "导入":
            continue
        if item.party_a_name:
            existed = _merge_party_a(item)
            stats["duplicates" if existed else "imported_units"] += 1
        if item.party_b_name:
            existed = _merge_party_b(item)
            stats["duplicates" if existed else "imported_units"] += 1
        if item.party_a_name and item.party_a_contact:
            existed = _merge_contact(item.party_a_name, item.party_a_contact, item.party_a_phone, item.party_a_email, item.party_a_mailing_address, seller=False)
            stats["duplicates" if existed else "imported_contacts"] += 1
        if item.party_b_name and item.party_b_contact:
            existed = _merge_contact(item.party_b_name, item.party_b_contact, item.party_b_phone, item.party_b_email, item.party_b_mailing_address, seller=True)
            stats["duplicates" if existed else "imported_contacts"] += 1
    return stats


def _merge_fill_empty(existing, incoming):
    merged = dict(existing)
    changed = False
    for key, value in incoming.items():
        if value and not str(merged.get(key, "") or "").strip():
            merged[key] = value
            changed = True
    return merged, changed


def _merge_party_a(item):
    data = read_json(config.F_PRIVATE_CUSTOMERS, [])
    data = data if isinstance(data, list) else list(data.values()) if isinstance(data, dict) else []
    incoming = {
        "PartyAName": item.party_a_name,
        "PartyAAddressPhone": item.party_a_address,
        "PartyACreditCode": item.party_a_credit,
        "PartyAAccountName": item.party_a_name,
        "PartyABank": item.party_a_bank,
        "PartyAAccountNo": item.party_a_account,
        "PartyAAddress": item.party_a_address,
        "PartyAContact": item.party_a_contact,
        "PartyAPhone": item.party_a_phone,
        "PartyAEmail": item.party_a_email,
        "UpdatedAt": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "SourceFiles": [item.path],
    }
    for index, existing in enumerate(data):
        name = existing.get("PartyAName") or existing.get("party_a_name") or ""
        if name == item.party_a_name:
            merged, _changed = _merge_fill_empty(existing, incoming)
            source_files = set(existing.get("SourceFiles") or [])
            source_files.add(item.path)
            merged["SourceFiles"] = sorted(source_files)
            data[index] = merged
            write_json(config.F_PRIVATE_CUSTOMERS, data)
            return True
    data.append(incoming)
    write_json(config.F_PRIVATE_CUSTOMERS, data)
    return False


def _merge_party_b(item):
    data = read_json(config.F_PARTY_B_LIBRARY, {})
    data = data if isinstance(data, dict) else {}
    incoming = {
        "name": item.party_b_name,
        "address": item.party_b_address,
        "credit": item.party_b_credit,
        "bank": item.party_b_bank,
        "account": item.party_b_account,
    }
    if item.party_b_name in data:
        merged, _changed = _merge_fill_empty(data[item.party_b_name], incoming)
        merged["name"] = item.party_b_name
        data[item.party_b_name] = merged
        write_json(config.F_PARTY_B_LIBRARY, data)
        return True
    data[item.party_b_name] = incoming
    write_json(config.F_PARTY_B_LIBRARY, data)
    return False


def _merge_contact(unit_name, contact, phone, email, mailing_address, seller):
    path = getattr(config, "F_SELLER_CONTACTS", "") if seller else config.F_PRIVATE_CONTACTS
    data = read_json(path, [])
    data = data if isinstance(data, list) else []
    incoming = {
        "unit_name": unit_name,
        "contact": contact,
        "phone": phone,
        "email": email,
        "mailing_address": mailing_address,
        "source_files": [],
        "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    for index, existing in enumerate(data):
        if existing.get("unit_name", "") == unit_name and existing.get("contact", "") == contact:
            merged, _changed = _merge_fill_empty(existing, incoming)
            data[index] = merged
            write_json(path, data)
            return True
    data.append(incoming)
    write_json(path, data)
    return False


def save_scan_log(summary, import_stats=None):
    log_path = Path(config.USER_DATA_DIR) / "scan_contract_log.txt"
    payload = {
        "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "summary": {
            "scanned_count": summary.get("scanned_count", 0),
            "success_count": summary.get("success_count", 0),
            "pending_count": summary.get("pending_count", 0),
            "low_confidence_count": summary.get("low_confidence_count", 0),
            "skipped_count": summary.get("skipped_count", 0),
            "unrecognized_count": summary.get("unrecognized_count", 0),
        },
        "files": [asdict(item) for item in summary.get("results", [])],
        "skipped": summary.get("skipped", []),
        "unrecognized": summary.get("unrecognized", []),
        "import": import_stats or {},
    }
    log_path.parent.mkdir(parents=True, exist_ok=True)
    log_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(log_path)


def save_scan_debug(summary):
    debug_path = Path(config.USER_DATA_DIR) / "scan_contract_debug.txt"
    lines = [f"扫描时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ""]
    for item in summary.get("results", []):
        lines.append("=" * 60)
        lines.append(f"扫描文件名：{item.file_name}")
        lines.append(f"文件路径：{item.path}")
        lines.append("最终写入字段：")
        lines.append(json.dumps(asdict(item), ensure_ascii=False, indent=2))
        if item.pending_info:
            lines.append("未确认信息：")
            lines.append(item.pending_info)
        if item.debug_notes:
            lines.append("字段提取明细：")
            lines.extend(item.debug_notes)
        lines.append("")
    if summary.get("skipped"):
        lines.append("跳过文件：")
        for path, reason in summary.get("skipped", []):
            lines.append(f"- {path}: {reason}")
    if summary.get("unrecognized"):
        lines.append("未识别文件：")
        for path, reason in summary.get("unrecognized", []):
            lines.append(f"- {path}: {reason}")
    debug_path.parent.mkdir(parents=True, exist_ok=True)
    debug_path.write_text("\n".join(lines), encoding="utf-8")
    return str(debug_path)
